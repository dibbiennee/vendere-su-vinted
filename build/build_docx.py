#!/usr/bin/env python3
"""Costruisce GUIDA_COMPLETA_VINTED_2026.docx dai markdown in capitoli/."""
import re, glob, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
CAPS = os.path.join(BASE, "capitoli")
OUT  = os.path.join(os.path.dirname(BASE), "GUIDA_COMPLETA_VINTED_2026.docx")

ACCENT   = RGBColor(0x0B, 0x76, 0x80)
ACCENT2  = RGBColor(0x11, 0x4B, 0x52)
GREY     = RGBColor(0x55, 0x55, 0x55)
BODYFONT = "Georgia"
HEADFONT = "Trebuchet MS"

# ---------------------------------------------------------------- helpers xml
def shade(el, hexcolor):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexcolor)
    el.append(sh)

def p_border(p, side='left', size=18, color='0B7680'):
    pPr = p._p.get_or_add_pPr()
    bd = pPr.find(qn('w:pBdr'))
    if bd is None:
        bd = OxmlElement('w:pBdr'); pPr.append(bd)
    e = OxmlElement('w:' + side)
    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(size))
    e.set(qn('w:space'), '8'); e.set(qn('w:color'), color)
    bd.append(e)

def p_shade(p, hexcolor):
    shade(p._p.get_or_add_pPr(), hexcolor)

def cell_shade(cell, hexcolor):
    shade(cell._tc.get_or_add_tcPr(), hexcolor)

def keep_with_next(p):
    p._p.get_or_add_pPr().append(OxmlElement('w:keepNext'))

def add_field(paragraph, instr):
    r = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    t  = OxmlElement('w:t'); t.text = "Aggiorna il campo per generare l'indice (Ctrl+A, poi F9)"
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for e in (f1, it, f2, t, f3):
        r._r.append(e)
    return r

def force_update_fields(doc):
    st = doc.settings.element
    uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true')
    st.append(uf)

# ---------------------------------------------------------------- doc setup
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(15.24), Cm(22.86)   # 6x9in, formato libro
sec.left_margin = sec.right_margin = Cm(1.9)
sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)

st = doc.styles
n = st['Normal']
n.font.name = BODYFONT; n.font.size = Pt(10)
n.element.rPr.rFonts.set(qn('w:eastAsia'), BODYFONT)
pf = n.paragraph_format
pf.line_spacing = 1.28; pf.space_after = Pt(7); pf.space_before = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def styleit(name, size, color, bold=True, before=0, after=6, font=HEADFONT, caps=False):
    s = st[name]
    s.font.name = font; s.font.size = Pt(size); s.font.bold = bold; s.font.color.rgb = color
    s.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after  = Pt(after)
    s.paragraph_format.line_spacing = 1.1
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if caps: s.font.all_caps = True

styleit('Heading 1', 21, ACCENT2, before=0,  after=6)
styleit('Heading 2', 13.5, ACCENT, before=16, after=5)
styleit('Heading 3', 11, ACCENT2, before=12, after=4)
st['Heading 1'].paragraph_format.keep_with_next = True
st['Heading 2'].paragraph_format.keep_with_next = True
st['Heading 3'].paragraph_format.keep_with_next = True

for nm in ('List Bullet', 'List Number'):
    s = st[nm]
    s.font.name = BODYFONT; s.font.size = Pt(10)
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.line_spacing = 1.22
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ---------------------------------------------------------------- inline md
INLINE = re.compile(r'(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)')

TOC_ENTRIES = []          # (livello, testo)
import json
PAGEMAP = {}
if os.path.exists(os.path.join(BASE, 'pagemap.json')):
    PAGEMAP = json.load(open(os.path.join(BASE, 'pagemap.json'), encoding='utf-8'))

def add_inline(p, text, base_size=None, base_color=None, italic_all=False):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(8.5)
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1]); r.italic = True
        else:
            r = p.add_run(part)
        if base_size: r.font.size = base_size
        if base_color is not None: r.font.color.rgb = base_color
        if italic_all: r.italic = True
    return p

# ---------------------------------------------------------------- copertina
COVER = os.path.join(BASE, 'cover.png')

def add_body_footer(section):
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(fp, ' PAGE ')
    for r in fp.runs:
        r.font.size = Pt(8.5); r.font.name = HEADFONT; r.font.color.rgb = GREY
    fp.runs[-1].text = ''

def cover(title, subtitle, edition):
    """Prima pagina: immagine a pieno formato, senza margini e senza numero."""
    s0 = doc.sections[0]
    s0.left_margin = s0.right_margin = s0.top_margin = s0.bottom_margin = Cm(0)
    s0.header_distance = s0.footer_distance = Cm(0)
    p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
    p.add_run().add_picture(COVER, width=s0.page_width)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width, body.page_height = Cm(15.24), Cm(22.86)
    body.left_margin = body.right_margin = Cm(1.9)
    body.top_margin = Cm(2.0); body.bottom_margin = Cm(2.0)
    body.header_distance = Cm(1.2); body.footer_distance = Cm(1.0)
    add_body_footer(body)

def toc_entry(level, text):
    """Riga d'indice con tabulazione a punti e numero di pagina."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2 if level == 2 else 5)
    p.paragraph_format.space_before = Pt(7 if level == 1 else 0)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.7 if level == 2 else 0)
    tabs = p.paragraph_format.tab_stops
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    tabs.add_tab_stop(Cm(11.4), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(text)
    if level == 1:
        r.font.name = HEADFONT; r.font.size = Pt(10); r.bold = True; r.font.color.rgb = ACCENT2
    else:
        r.font.name = BODYFONT; r.font.size = Pt(9); r.font.color.rgb = GREY
    r2 = p.add_run('\t' + str(PAGEMAP.get(text, '00')))
    r2.font.name = HEADFONT; r2.font.size = Pt(9)
    r2.font.color.rgb = ACCENT2 if level == 1 else GREY
    if level == 1: r2.bold = True

def toc_page():
    p = doc.add_paragraph()
    r = p.add_run('Indice')
    r.font.name = HEADFONT; r.font.size = Pt(21); r.font.bold = True; r.font.color.rgb = ACCENT2
    p.paragraph_format.space_after = Pt(4)
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(0); hr.paragraph_format.space_after = Pt(13)
    hr.paragraph_format.line_spacing = 1.0
    hr.add_run().font.size = Pt(2)
    p_border(hr, 'bottom', 10, '0B7680')
    for lvl, txt in TOC_ENTRIES:
        toc_entry(lvl, txt)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ---------------------------------------------------------------- blocchi
def add_quote(lines):
    txt = ' '.join(lines).strip()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_border(p, 'left', 24, '0B7680')
    p_shade(p, 'F2F8F8')
    add_inline(p, txt, base_size=Pt(10))

def add_code(lines):
    for i, ln in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(6) if i == 0 else Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_shade(p, 'F5F5F3')
        r = p.add_run(ln if ln.strip() else ' ')
        r.font.name = 'Consolas'; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x33,0x33,0x33)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)

def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true')
    trPr.append(th)

def no_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))

def add_table(rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.05
        add_inline(p, h, base_size=Pt(8.5))
        for r in p.runs: r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.name = HEADFONT
        cell_shade(c, '0B7680')
    repeat_header(t.rows[0]); no_split(t.rows[0])
    for bi, row in enumerate(body):
        cells = t.add_row().cells
        for i, val in enumerate(row[:len(header)]):
            c = cells[i]; c.text = ''
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.08
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, val, base_size=Pt(8.5))
            for r in p.runs: r.font.name = BODYFONT
            if bi % 2 == 1: cell_shade(c, 'F4F7F7')
        no_split(t.rows[-1])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(10)
    p_border(p, 'bottom', 6, 'CCCCCC')

# ---------------------------------------------------------------- parser
def render(md, first=False):
    lines = md.split('\n')
    i = 0
    tbuf, cbuf, qbuf = [], [], []
    incode = False
    title = sub = ed = None

    def flush_table():
        nonlocal tbuf
        if tbuf: add_table(tbuf); tbuf = []
    def flush_quote():
        nonlocal qbuf
        if qbuf: add_quote(qbuf); qbuf = []

    while i < len(lines):
        ln = lines[i]; s = ln.strip()

        if incode:
            if s.startswith('```'):
                add_code(cbuf); cbuf = []; incode = False
            else:
                cbuf.append(ln)
            i += 1; continue

        if s.startswith('%TITLE%'):   title = s[7:].strip(); i += 1; continue
        if s.startswith('%SUBTITLE%'):sub   = s[10:].strip(); i += 1; continue
        if s.startswith('%EDITION%'):
            ed = s[9:].strip()
            cover(title, sub, ed); toc_page(); i += 1; continue

        if s.startswith('```'):
            flush_table(); flush_quote(); incode = True; i += 1; continue

        if s.startswith('|') and s.endswith('|'):
            cells = [c.strip() for c in s.strip('|').split('|')]
            if not all(re.fullmatch(r':?-{2,}:?', c) for c in cells):
                tbuf.append(cells)
            i += 1; continue
        else:
            flush_table()

        if s.startswith('>'):
            qbuf.append(s.lstrip('>').strip()); i += 1; continue
        else:
            flush_quote()

        if not s:
            i += 1; continue

        if s == '---':
            add_hr(); i += 1; continue

        if s.startswith('# '):
            if not first or doc.paragraphs[-1].text.strip():
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            p = doc.add_paragraph(style='Heading 1')
            add_inline(p, s[2:])
            hr = doc.add_paragraph()
            hr.paragraph_format.space_before = Pt(0); hr.paragraph_format.space_after = Pt(13)
            hr.paragraph_format.line_spacing = 1.0
            hr.add_run().font.size = Pt(2)
            p_border(hr, 'bottom', 10, '0B7680')
            i += 1; continue

        if s.startswith('## '):
            p = doc.add_paragraph(style='Heading 2')
            add_inline(p, s[3:]); i += 1; continue
        if s.startswith('### '):
            p = doc.add_paragraph(style='Heading 3'); add_inline(p, s[4:]); i += 1; continue

        m = re.match(r'^- \[ \] (.*)', s)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.6)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run('☐  '); r.font.size = Pt(11); r.font.color.rgb = ACCENT
            add_inline(p, m.group(1))
            i += 1; continue

        m = re.match(r'^[-*] (.*)', s)
        if m:
            p = doc.add_paragraph(style='List Bullet'); add_inline(p, m.group(1)); i += 1; continue

        m = re.match(r'^(\d+)\. (.*)', s)
        if m:
            p = doc.add_paragraph(style='List Number'); add_inline(p, m.group(2)); i += 1; continue

        p = doc.add_paragraph(); add_inline(p, s)
        i += 1

    flush_table(); flush_quote()

files = sorted(glob.glob(os.path.join(CAPS, '*.md')))
SOURCES = []
for f in files:
    with open(f, encoding='utf-8') as fh:
        SOURCES.append(fh.read())

# pre-scansione: raccoglie le voci d'indice prima di costruire la pagina Indice
for src in SOURCES:
    for ln in src.split('\n'):
        t = ln.strip()
        if t.startswith('# '):
            TOC_ENTRIES.append((1, re.sub(r'\*\*|`', '', t[2:])))
        elif t.startswith('## '):
            TOC_ENTRIES.append((2, re.sub(r'\*\*|`', '', t[3:])))
PRESCANNED = list(TOC_ENTRIES)
TOC_ENTRIES.clear(); TOC_ENTRIES.extend(PRESCANNED)
_seen_toc = {'done': False}

for idx, src in enumerate(SOURCES):
    render(src, first=(idx == 0))

cp = doc.core_properties
cp.title = 'Vendere su Vinted — La guida completa'
cp.subject = 'Guida operativa alla vendita su Vinted'
cp.comments = 'Edizione 2026'
cp.author = ''
cp.last_modified_by = ''
doc.save(OUT)
print("scritto:", OUT)
